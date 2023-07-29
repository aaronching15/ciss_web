# -*- encoding:"utf-8"-*-
__author__ = " ruoyu.Cheng"

'''
=============================================== 
todo:
功能：ciss_db对应db_funda数据库地数据读写
def fund_data_manage(request): ### 基金净值、收益率等数据管理
def fund_analysis 基金经理定性分析数据表，实时维护
def fund_data_manage  基金净值、收益率等数据管理

def fund_analysis_old 基金定性分析功能，2022-9之前的旧版本
# 202303开始，fund_analysis_old 功能都迁移至 fund_data_manage

last 20230315 | since 20220823
refernce:derived from views_pms_manage.py
===============================================
''' 

###########################################################################
### Initialization
from operator import index
from django.shortcuts import render
from django.http import JsonResponse
from django.http import HttpResponse 
from django.views.decorators.csrf import csrf_protect,requires_csrf_token,csrf_exempt

###########################################################################
### 导入项目地址
import sys,os

# 当前目录 C:\zd_zxjtzq\ciss_web\CISS_rc\db
path_ciss_web = os.getcwd().split("CISS_rc")[0]
path_ciss_rc = path_ciss_web +"\\CISS_rc\\"
sys.path.append(path_ciss_rc + "config\\" )
sys.path.append(path_ciss_rc + "db\\" )
sys.path.append(path_ciss_rc + "db\\fund_analysis\\" )
sys.path.append(path_ciss_rc + "db\\db_assets\\" )
sys.path.append(path_ciss_rc + "db\\analysis_indicators\\" )
sys.path.append(path_ciss_rc + "db\\data_io\\" )

import pandas as pd
import numpy as np 
import json 
sys.path.append("..") 
import datetime as dt  
time_now = dt.datetime.now()
### 获取近1w、1m、3m、6m、1year日期
time_now_str = dt.datetime.strftime(time_now   , "%Y%m%d")
time_now_str_pre1d = dt.datetime.strftime(time_now - dt.timedelta(days=1)  , "%Y%m%d")

######################################################################################################
######################################################################################################
### fund_analysis_old 旧版基金分析，202303开始，fund_analysis_old 功能都迁移至 fund_data_manage
# def fund_analysis_old(request):
#     ###########################################################################
#     ### 旧版基金分析，202303开始，fund_analysis_old 功能都迁移至 fund_data_manage 
#     return render(request,"ciss_exhi/fund_fof/index_fund_fof.html",context) 

    
######################################################################################################
######################################################################################################
@csrf_protect
@requires_csrf_token
@csrf_exempt
def fund_analysis(request): 
    #################################################################################
    ### 定义dict对象 context
    context={"info":"none"} 
    ### 时间
    # context["time_now_str"] = time_now_str
    # context["time_now_str_pre1d"] =time_now_str_pre1d 

    return render(request,"ciss_exhi/fund_fof/fund_analysis.html",context) 


######################################################################################################
# GETDATA 获取table数据 
@csrf_protect
@requires_csrf_token
@csrf_exempt
def get_data(request):
    ### layui的table设置里，默认是 get
    if request.method == 'GET':
        print("get_data | request.GET", type(request.GET.keys() ),request.GET.keys() )
        print("get_data | request.POST", type(request.POST.keys() ),request.POST.keys() )
        #########################################################
        ### 
        #######################################################
        ### 查找的input变量        
        dict_select = {}
        col_list = ["id","date","name","code","style_fund","theme_fund","ind_1","ind_2","ind_3","ind_num"]
        col_list = col_list + ["score_performance","s_down_market","s_flat_market","s_up_market","abstract_analysis","fund_manager"]
        for temp_col in col_list : 
            if temp_col in request.GET.keys() :
                if len( request.GET[temp_col] ) > 0 :
                    dict_select[temp_col] = request.GET[temp_col]
        print("Debug | dict_select=", dict_select ) 
        
        #######################################################
        ### 获取表格数据
        obj_db = {} 
        obj_db["db_name"] = "db_funda.sqlite3"
        obj_db["table_name"] = "ciss_exhi"+ "_" + "fund_analysis" 
        obj_db["dict_select"] = dict_select 
        
        #############################################
        from database import db_sqlite
        db_sqlite1 = db_sqlite()
        obj_db = db_sqlite1.select_table_data( obj_db ) 
        df_data = obj_db["df_data"]  
        ### 
        df_data = df_data.sort_values(by="id" ,ascending=False )
        # print("Debug:End of get_data")
        
        #############################################
        ### 数据格式转换：df转换成layui可识别的dict
        ### notes:df不能直接df.to_json(),这样得到的是 str格式，df_data.to_json()不能赋值给JsonResponse里的data
        ### layui-table可以识别的data数据在views.py里长这样：dict_data= {1: {'id': 1, 'date': '600036','weight': 0.05},  
        # 2: {'id': 2, 'date': '300750', 'weight': 0.08},  
        col_list_output= col_list
        dict_data = {}
        for i in df_data.index :
            dict_data[i]={}
            for temp_col in col_list_output : 
                try :
                    dict_data[i][temp_col] = df_data.loc[i, temp_col] 
                except:
                    print("Debug: temp_col",temp_col,i ,"\n", df_data.loc[i, :] ) 
                    asd

        
        ### notes: 必须用 df_data.T.to_dict() ,如果 dict( df_data ) 会报错
        ### notes：如果数据库记录里有多列是null空值|例如测试用的列，就会导致这个问题！！
        # dict_data = df_data.loc[ -3: ,col_list ].T.to_dict() 
        dict_data = df_data.loc[ : ,col_list ].T.to_dict() 
        count = len( dict_data )
        print("Debug-get_data：type of data=",type(dict_data),count  )  

        ##############################################        
        ### table 组件默认规定的数据格式：{ "code": 0, "msg": "",  "count": 1000,  "data": [{}, {}] } 
        ### 非默认的格式需要parseData转换，例如 {"status": 0,   "message": "", "total": 180, "data": {"item": [{}, {}] } }
        ### 官方格式参考，url=http://layui.winxapp.cn/demo/table/user/-page=1&limit=30.js
        if count != 0:
            # print("count=",count,"dict_data=", dict_data[0] )
            return JsonResponse({'code':0,'msg':'查询成功','count':count,'data':dict_data })
            # return JsonResponse({'code':0,'msg':'查询成功','count':count,'data':json_data })
        else:
            return JsonResponse({'code':0,'msg':'暂无数据','count':count,'data':{}  })





######################################################################################################
### 新增记录
@csrf_protect
@requires_csrf_token
@csrf_exempt
def add_data(request):
    ###  
    print("add_data | request.GET", type(request.GET.keys() ),request.GET.keys() )
    print("add_data | request.POST", type(request.POST.keys() ),request.POST.keys() ) 
    if request.method == 'POST': 
        print("add_data | request.POST", type(request.POST.keys() ),request.POST.keys() ) 
        ##################################################################################
        ### 存入sql数据表 
        ### 必须项 ["id","date","name","code","style_fund","theme_fund","ind_1","ind_2","ind_3","ind_num"]
        id = request.POST.get("id","")
        date = request.POST.get("date","") 
        code = request.POST.get("code","") 
        name = request.POST.get("name","") 
        style_fund = request.POST.get("style_fund","")
        theme_fund = request.POST.get("theme_fund","") 
        ind_1 = request.POST.get("ind_1","") 
        ind_2 = request.POST.get("ind_2","") 
        ind_3 = request.POST.get("ind_3","") 
        ind_num = request.POST.get("ind_num","") 
        ### 必须项 ["score_performance","s_down_market","s_flat_market","s_up_market","abstract_analysis","fund_manager"]
        score_performance = request.POST.get("score_performance","") 
        s_down_market = request.POST.get("s_down_market","") 
        s_flat_market = request.POST.get("s_flat_market","") 
        s_up_market = request.POST.get("s_up_market","")
        abstract_analysis = request.POST.get("abstract_analysis","") 
        fund_manager = request.POST.get("fund_manager","")  
        ### 计算综合得分
        if len(score_performance) > 0 : 
            if not float(score_performance) > 1 :
                ###
                try :
                    temp = float(s_down_market) +float(s_flat_market) + float(s_up_market)
                    score_performance = str( round( temp,2 ) )
                except :
                    pass
        else :
            score_performance = "0.0"
        # print("Debug: date=",date, )
        ### 非必须项 ["note","date_lastmodify","if_fundmanager_fault"]
        note = request.POST.get("note","")
        date_lastmodify = request.POST.get("date_lastmodify","") 
        if_fundmanager_fault = request.POST.get("if_fundmanager_fault","") 
        
        print("Debug: add data=",date,";code=",code, ";name=",code,"abstract_analysis=", abstract_analysis)
        ##################################################################################
        ### 必须项不能是空的
        if date == '' or code == '' : 
            return JsonResponse({'code': 10021, 'msg': '参数错误'}, json_dumps_params={'ensure_ascii': False})
        else :
            ###############################################
            ### 多个input必须至少有一项非空
            obj_db = {} 
            obj_db["db_name"] = "db_funda.sqlite3"
            obj_db["table_name"] = "ciss_exhi"+ "_" + "fund_analysis" 
            obj_db["insert_type"] = "1r"
            obj_db["dict_1r"] = {}
            if len( date ) > 0  :  
                obj_db["dict_1r"][ "date"] = date 
            if len( code ) > 0  :  
                obj_db["dict_1r"][ "code"] = code  
            if len( name ) > 0  :  
                obj_db["dict_1r"][ "name"] = name
            if len( style_fund ) > 0  :  
                obj_db["dict_1r"][ "style_fund"] = style_fund 
            if len( theme_fund ) > 0  :  
                obj_db["dict_1r"][ "theme_fund"] = theme_fund  
            if len( ind_1 ) > 0  :  
                obj_db["dict_1r"][ "ind_1"] = ind_1
            if len( ind_2 ) > 0  :  
                obj_db["dict_1r"][ "ind_2"] = ind_2
            if len( ind_3 ) > 0  :  
                obj_db["dict_1r"][ "ind_3"] = ind_3
            if len( ind_num ) > 0  :  
                obj_db["dict_1r"][ "ind_num"] = ind_num 
            if len( score_performance ) > 0  :  
                obj_db["dict_1r"][ "score_performance"] = score_performance  
            if len( s_down_market ) > 0  :  
                obj_db["dict_1r"][ "s_down_market"] = s_down_market
            if len( s_flat_market ) > 0  :  
                obj_db["dict_1r"][ "s_flat_market"] = s_flat_market 
            if len( s_up_market ) > 0  :  
                obj_db["dict_1r"][ "s_up_market"] = s_up_market  
            if len( abstract_analysis ) > 0  :  
                obj_db["dict_1r"][ "abstract_analysis"] = abstract_analysis
            if len( fund_manager ) > 0  :  
                obj_db["dict_1r"][ "fund_manager"] = fund_manager 
            if len( date_lastmodify ) > 0  :  
                obj_db["dict_1r"][ "date_lastmodify"] = date_lastmodify
            if len( if_fundmanager_fault ) > 0  :  
                obj_db["dict_1r"][ "if_fundmanager_fault"] = if_fundmanager_fault 
            if len( note ) > 0  :  
                obj_db["dict_1r"][ "note"] = note 

            #############################################
            from database import db_sqlite
            db_sqlite1 = db_sqlite()
            obj_db = db_sqlite1.insert_table_data(obj_db) 

            return JsonResponse({'code': 0, 'msg': '添加成功！'})

    else :
        ### 
        print("Not post for add data")


######################################################################################################
### 编辑数据表记录
@csrf_protect
@requires_csrf_token
@csrf_exempt
def edit_data(request):
    if request.method == 'POST':
        ### request.POST <class 'dict_keys'> dict_keys(['id', 'date', 'strategy_CN', 'code', 'name', 'weight'])
        print("edit_data |request.POST", type(request.POST.keys() ),request.POST.keys() )
        
        #########################################################
        ### html里提交修改后的数据,传给后台view的变量是formData1.append("id")里的 id，不是edit_id
        id = request.POST.get('id',"")
        date= request.POST.get('date',"")
        name= request.POST.get('name',"")
        code= request.POST.get('code',"")
        style_fund= request.POST.get('style_fund',"")
        theme_fund= request.POST.get('theme_fund',"")
        ind_1= request.POST.get('ind_1',"")
        ind_2= request.POST.get('ind_2',"")
        ind_3= request.POST.get('ind_3',"")
        ind_num= request.POST.get('ind_num',"")#
        score_performance= request.POST.get('score_performance',"")
        s_down_market= request.POST.get('s_down_market',"")
        s_flat_market= request.POST.get('s_flat_market',"")
        s_up_market= request.POST.get('s_up_market',"")
        abstract_analysis= request.POST.get('abstract_analysis',"")
        fund_manager= request.POST.get('fund_manager',"")#
        # 非必须项
        note= request.POST.get('note',"")
        date_lastmodify= request.POST.get('date_lastmodify',"")
        if_fundmanager_fault= request.POST.get('if_fundmanager_fault',"") 

        ### 
        ###############################################
        ### 多个input必须至少有一项非空
        obj_db = {} 
        obj_db["db_name"] = "db_funda.sqlite3"
        obj_db["table_name"] = "ciss_exhi"+ "_" + "fund_analysis" 
        obj_db["insert_type"] = "1r"
        obj_db["dict_1r"] = {}
        ## id是必备项
        obj_db["dict_1r"]["id"]  = id 
        if len( date ) > 0  :  
            obj_db["dict_1r"][ "date"] = date
                    
        if len( name ) > 0  :  
            obj_db["dict_1r"]["name"] = name
        if len( code ) > 0  :  
            obj_db["dict_1r"][ "code"] = code
        if len( style_fund ) > 0  :  
            obj_db["dict_1r"]["style_fund"] =style_fund
        if len( theme_fund ) > 0  :  
            obj_db["dict_1r"]["theme_fund"]  = theme_fund
        if len( ind_1 ) > 0  :  
            obj_db["dict_1r"]["ind_1"] =ind_1
        if len( ind_2 ) > 0  :  
            obj_db["dict_1r"]["ind_2"] =ind_2
        if len( ind_3 ) > 0  :  
            obj_db["dict_1r"]["ind_3"] =ind_3
        if len( ind_num ) > 0  :  
            obj_db["dict_1r"][ "ind_num"] = ind_num 
        if len( score_performance ) > 0  :  
            obj_db["dict_1r"][ "score_performance"] = score_performance
        if len( s_down_market ) > 0  :  
            obj_db["dict_1r"][ "s_down_market"] = s_down_market
        if len( s_flat_market ) > 0  :  
            obj_db["dict_1r"][ "s_flat_market"] = s_flat_market
        if len( s_up_market ) > 0  :  
            obj_db["dict_1r"][ "s_up_market"] = s_up_market
        if len( abstract_analysis ) > 0  :  
            obj_db["dict_1r"][ "abstract_analysis"] = abstract_analysis
        if len( fund_manager ) > 0  :  
            obj_db["dict_1r"][ "fund_manager"] = fund_manager
        if len( note ) > 0  :  
            obj_db["dict_1r"][ "note"] = note 
        if len( date_lastmodify ) > 0  :  
            obj_db["dict_1r"][ "date_lastmodify"] = date_lastmodify
        if len( if_fundmanager_fault ) > 0  :  
            obj_db["dict_1r"][ "if_fundmanager_fault"] = if_fundmanager_fault 
        # print("Debug=dict_1r=", obj_db["dict_1r"] )
        #############################################
        from database import db_sqlite
        db_sqlite1 = db_sqlite()
        ### "update_type" = "id" ：对某一列id的值进行更新
        obj_db["update_type"] = "id"
        obj_db["id"] = id
        
        try:
            obj_db = db_sqlite1.update_table_data(obj_db) 
            return JsonResponse({'code':0,'msg':'修改成功!'})
        except :
            return JsonResponse({'status': 10023, 'msg': '数据不存在'})
            # return JsonResponse({'code':0, 'msg': '数据不存在'})




######################################################################################################
### 删除数据表记录
@csrf_protect
@requires_csrf_token
@csrf_exempt
def del_data(request):
    if request.method == 'POST':
        print("request.POST", type(request.POST.keys() ),request.POST.keys() )
        #########################################################
        ### 根据提交的ID删除记录
        id_table = request.POST.get('id',"")

        obj_db = {} 
        ### 数据库信息
        obj_db["db_name"] = "db_funda.sqlite3"
        obj_db["table_name"] = "ciss_exhi"+ "_" + "fund_analysis" 
        ### obj_db["delete_type"] ==  "id"  根据给定id 删除某一行
        obj_db["delete_type"] = "id" 
        obj_db["id_table"] = id_table
        #############################################

        try :
            from database import db_sqlite
            db_sqlite1 = db_sqlite()     
            obj_db = db_sqlite1.delete_table_index(obj_db) 
            ### 无返回数据 
            return JsonResponse({'code': 0, 'msg': '删除成功','count':id_table } )

        except :
            return JsonResponse({'code':0,'msg':'删除过程出错','count':id_table } )



######################################################################################################
######################################################################################################
@csrf_protect
@requires_csrf_token
@csrf_exempt
def fund_data_manage(request): 
    #################################################################################
    ### 基金净值、收益率等数据管理
    ### 定义dict对象 context
    print("request.POST", type(request.POST.keys() ),request.POST.keys() )
    context={"info":"none"} 
    ### 时间
    context["time_now_str"] = time_now_str
    context["time_now_str_pre1d"] =time_now_str_pre1d  
    
    ### 导入常用地址信息
    import os
    path_pms =  os.getcwd().split("ciss_web")[0]+"\\data_pms\\"  
    path_wind_terminal = path_pms + "wind_terminal\\" 
    path_wss = path_pms + "wss\\"
    path_wpf = path_pms + "wpf\\"
    path_wpd = path_pms + "wpd\\"
    path_wsd = path_pms + "wsd\\"
    path_ciss_web = os.getcwd().split("ciss_web")[0]+"ciss_web\\"
    path_ciss_rc = path_ciss_web +"CISS_rc\\"
    # path_db = path_ciss_rc + "db\\"
    path_dt = path_ciss_rc + "db\\db_times\\"
    sys.path.append(path_ciss_rc+ "config\\")
    from config_data import config_db
    config_db_1 = config_db()
    path_db = path_ciss_web

    #################################################################################
    if "input_fund_unit" in request.POST.keys():     
        from funds import fund_exhi
        class_fund_exhi = fund_exhi()    
        ### 
        obj_f ={} 
        ### obj_f[ "exhi_type"] = "unit"对应单位净值 ; or "perf"
        obj_f[ "exhi_type"] = "unit"
        obj_f["fund_code"]= request.POST.get("fund_code","") 
        ### 获取区间开始和结束日期
        obj_f[ "date_begin"]= request.POST.get("date_begin","20201231")
        obj_f[ "date_end"]  = request.POST.get("date_end","20211231") 
        ### 提取的指标
        col_list_wind_api = request.POST.get("col_list_wind_api","") 
        if col_list_wind_api != "NAV_adj" and len(col_list_wind_api )> 5 :
            ### 可能有多个指标
            obj_f["col_list_wind_api"]= col_list_wind_api

        ##########################################
        ### 判断是否需要提取基准指数 ；偏股混合型基金指数 885001.WI
        temp_bench_code = request.POST.get("bench_code","") 
        if len(  temp_bench_code ) > 0 : 
            obj_f[ "bench_code"]= temp_bench_code
        
        ### 
        obj_f = class_fund_exhi.get_fund_exhi(obj_f) 
        ### 需要转置transpose
        context["df_unit"] =  obj_f["df_unit"].T  
        if "date_hist_start" in obj_f.keys() :
            context["date_hist_start"]  = obj_f["date_hist_start"] 
        if "date_hist_end" in obj_f.keys() :
            context["date_hist_end"]  = obj_f["date_hist_end"]  
        

    #################################################################################
    if "input_fund_unit_list" in request.POST.keys():     
        ##########################################
        ### 导入基金列表
        file_name = request.POST.get("file_name","") 
        file_path = request.POST.get("file_path","") 
        sheet_name1 = request.POST.get("sheet_name","") 
        col_name = request.POST.get("col_name","") 
        # C:\rc_2023\rc_202X\data_pms\wsd 
        print("file_name" , file_path +file_name, sheet_name1 , col_name)
        df_temp = pd.read_excel( file_path+ file_name, sheet_name=sheet_name1 )

        code_list = list( df_temp[ col_name ] )
        print("code_list :", code_list  )
        
        ##########################################
        ### 设置统一的变量
        from funds import fund_exhi
        class_fund_exhi = fund_exhi()    
        obj_f0 ={}             
        ### obj_f[ "exhi_type"] = "unit"对应单位净值 ; or "perf"
        obj_f0[ "exhi_type"] = "unit" 
        ### 获取区间开始和结束日期
        obj_f0[ "date_begin"]= request.POST.get("date_begin","20201231")
        obj_f0[ "date_end"]  = request.POST.get("date_end","20211231") 
        ### 提取的指标
        col_list_wind_api = request.POST.get("col_list_wind_api","") 
        if col_list_wind_api != "NAV_adj" and len(col_list_wind_api )> 5 :
            ### 可能有多个指标
            obj_f0["col_list_wind_api"]= col_list_wind_api
        ### 判断是否需要提取基准指数 ；偏股混合型基金指数 885001.WI
        temp_bench_code = request.POST.get("bench_code","") 
        if len(  temp_bench_code ) > 0 : 
            obj_f[ "bench_code"]= temp_bench_code

        ##########################################
        for temp_code in code_list :
            ### 
            obj_f= obj_f0
            obj_f["fund_code"]= temp_code
            obj_f = class_fund_exhi.get_fund_exhi(obj_f) 
            

        ##########################################
        ### 需要转置transpose
        context["df_unit"] =  obj_f["df_unit"].T      


    
    #################################################################################
    ### Case 1：筛选特定基金记录
    if "input_fund_select" in request.POST.keys():   
        ### 
        fund_data_source = request.POST.get("fund_data_source","") 
        ### 
        fund_code = request.POST.get("fund_search_fund_code","")
        if len( fund_code ) < 1 :        
            col_name = request.POST.get("fund_search_col_name","")
            col_value = request.POST.get("fund_search_col_value","")

            print("fund_code= " ,fund_code,"fund_data_source=",fund_data_source,col_name ,";col_name=",col_name , col_value ) 
        ################################################################################
        ### 直接拉取完成的基金表格，file= FF-基金研究-主动股票-220812.xlsx 
        if fund_data_source == "fund_search_quant_wind" : 
            ### 先统一导入所有基金代码
            ### 依次在 FF-基金研究-主动股票-220812.xlsx 里匹配
            ### path_root = 'C:\\rc_2023\\rc_202X\\'
            path_root = os.getcwd().split("ciss_web")[0]  
            path_data_pms = path_root +"data_pms\\"
            path_data_adj = path_root +"data_pms\\data_adj\\"

            path_fund_data = path_data_pms +"wind_terminal\\"
            fund_type_list = ["主动股票", "偏股混合", "偏债混合", "纯债", "股票指数", "可转债", "纯债", "FOF"]
            count_f = 0 
            for fund_type in fund_type_list :
                file_name = "FF-基金研究-" + fund_type +"-220812.xlsx"
                df_temp = pd.read_excel( path_fund_data+ file_name, sheet_name="概况" )
                ###
                if count_f == 0 :
                    df_f_all = df_temp
                    count_f = 1
                else :
                    df_f_all = df_f_all.append(df_temp )
            ### 匹配基金
            df_match = df_temp[ df_temp["基金代码"]== fund_code  ]
            if (df_match.index) > 0 :
                print( df_match.T )            
            else :
                print("No record in excel file ")
            

        ################################################################################
        ### 抓取数据库内匹配的信息
        if fund_data_source == "fund_search_db_funda" : 
            obj_db = {} 
            obj_db["db_name"] = "db_funda.sqlite3"
            obj_db["table_name"] = "ciss_exhi"+ "_" + "fund_analysis" 
            #############################################
            from database import db_sqlite
            db_sqlite1 = db_sqlite()
            #############################################
            ### 检查是否有基金代码输入，例如 110011.OF
            if len( fund_code ) >= 6 : 
                
                #############################################
                ### 有2个列的情况
                if len( col_name ) > 0 and len(col_value)>0   :
                    ### 需要匹配 2~N 列 |dict as input 
                    obj_db["dict_select"] = { }
                    obj_db["dict_select"]["code"] = fund_code
                    obj_db["dict_select"][ col_name ] = col_value         
                    obj_db = db_sqlite1.select_table_data( obj_db )
                    
                #############################################
                ### 只有1个列=基金代码的情况
                else :
                    obj_db["col_name"] = "code"
                    obj_db["select_value"] = fund_code
                    obj_db = db_sqlite1.select_table_data( obj_db )
                    
            else :
                #############################################
                ### 没有代码，只有1个列的情况
                ### 需要匹配 2~N 列      
                obj_db["dict_select"] = { }
                obj_db["dict_select"][col_name] = col_value 
                print("Debug 222= col_name " ,col_name  , col_value )
                print("dict_select \n", obj_db["dict_select"] )
                obj_db = db_sqlite1.select_table_data( obj_db )

        ################################################################################
        ### 返回数据库中的最新N条记录  
        df_data = obj_db["df_data"]
        df_data = df_data.sort_values(by="date", ascending=False)
        context["df_fund_research_tail"] = df_data.T 


    #################################################################################
    ### Case 2：获取基金定性研究记录，上传至sql数据库
    if "input_fund_research_to_sql" in request.POST.keys():    
        ############################################
        dict_1r ={}
        ############################################ dict_1r[""]
        ### 以下部分信息是必须的，有的如果len()<1,就提取Wind-api
        dict_1r["date"] = request.POST.get("date_lastmodify","")
        dict_1r["code"] = request.POST.get("fund_code","")
        dict_1r["style_fund"] = request.POST.get("style_fund","")
        dict_1r["theme_fund"] = request.POST.get("theme_fund","")
        dict_1r["ind_1"] = request.POST.get("ind_1","")
        dict_1r["ind_2"] = request.POST.get("ind_2","")
        dict_1r["ind_3"] = request.POST.get("ind_3","")
        dict_1r["ind_num"] = request.POST.get("ind_num","")
        dict_1r["s_down_market"] = request.POST.get("s_down_market","")
        dict_1r["s_flat_market"] = request.POST.get("s_flat_market","")
        dict_1r["s_up_market"] = request.POST.get("s_up_market","")
        dict_1r["abstract_analysis"]= request.POST.get("abstract_analysis","")

        ### 计算总分
        score_performance= float( dict_1r["s_down_market"] ) + float( dict_1r["s_flat_market"] ) +float( dict_1r["s_up_market"]  )
        score_performance= round( score_performance /30*100, 2 ) 
        dict_1r["score_performance"] = score_performance

        ############################################
        ### 以下部分信息不是必须的
        dict_1r["fund_manager"] = request.POST.get("fund_manager","")
        dict_1r["note"] = request.POST.get("note","")
        dict_1r["date_lastmodify"] = request.POST.get("date_lastmodify","")
        dict_1r["if_fundmanager_fault"] = request.POST.get("if_fundmanager_fault","0")


        ################################################################################
        ### 判断是否要匹配基金信息、或提取Wind-API指标
        if len( dict_1r["fund_manager"] ) <= 1 :
            obj_f={}
            obj_f["fund_code"] = dict_1r["code"] 
            obj_f["date"] = dict_1r["date_lastmodify"] 
            obj_f["col_list"] = ["fund_fundmanageroftradedate",]
            from get_wind_api import wind_api
            wind_api1 = wind_api()
            obj_f = wind_api1.get_wss_fund_1date( obj_f )

            ### obj_f["list_data"] = obj_w.Data[0] = ['张坤']
            dict_1r["fund_manager"] = obj_f["list_data"][0]

        ### 获取基金名称
        obj_f={}
        obj_f["fund_code"] = dict_1r["code"] 
        obj_f["date"] = dict_1r["date_lastmodify"] 
        obj_f["col_list"] = ["fund_info_name",]
        from get_wind_api import wind_api
        wind_api1 = wind_api()
        obj_f = wind_api1.get_wss_fund_1date( obj_f )

        ### obj_f["list_data"] = obj_w.Data[0] = ['张坤'] 
        dict_1r["name"] = obj_f["list_data"][0] 

        ################################################################################
        ### 将变量转换为sql语言，存入数据表 db_funda\ciss_exhi_fund_analysis
        from database import db_sqlite
        db_sqlite1 = db_sqlite()
        ### insert_table_data |  "1r" 导入一行
        obj_db = {}
        obj_db["db_name"] = "db_funda.sqlite3"
        obj_db["table_name"] =  "ciss_exhi"+ "_" + "fund_analysis"
        obj_db["insert_type"] ="1r"
        obj_db["dict_1r"] = dict_1r
        obj_db = db_sqlite1.insert_table_data(obj_db)

        ### 返回数据库中的最新N条记录 
        context["df_fund_research_tail"] = obj_db["df_data"].T 

        ################################################################################
        ### 返回输入项对应的 ID；如果输入有误，可以输入ID进行修改 TODO

    #################################################################################
    ### Case 3：删除基金某一行 | 
    ### 删除表:  drop删除表结构和表数据，truncate删除表数据，delete删除某一行
    if "input_fund_delete" in request.POST.keys():   
        fund_code= request.POST.get("del_fund_code","")
        ### 检查是否还需要匹配列
        col_name = request.POST.get("del_col_name","")
        col_value = request.POST.get("del_col_value","")
        obj_db = {}         
        obj_db["db_name"] = "db_funda.sqlite3"
        obj_db["table_name"] = "ciss_exhi"+ "_" + "fund_analysis" 
        #############################################
        from database import db_sqlite
        db_sqlite1 = db_sqlite()
        #############################################
        ### 有2个列的情况
        if len( col_name ) > 0 and len(col_value)>0   :
            ### 需要匹配 2~N 列               
            obj_db["col_list"] = ["code", col_name ]
            obj_db["value_list"] = [fund_code , col_value ]
            obj_db["delete_type"] = "1r" ###
            obj_db = db_sqlite1.delete_table( obj_db )

        #############################################
        ### 只有1个列=基金代码的情况
        else :
            input1 = input("Check if you want to delete all records match code=", fund_code )
            obj_db["col_name"] = "code"            
            obj_db["select_value"] = fund_code
            obj_db = db_sqlite1.delete_table( obj_db )

        ################################################################################
        ### 返回数据库中的最新N条记录 
        df_data = obj_db["df_data"].sort_values(by="id",ascending=False )
        df_data = df_data.head(10)

        context["df_fund_research_tail"] = df_data.T 




    ###########################################################################
    ### 基金池计算
    if "input_cal_fundpool" in request.POST.keys():  
        ### 基金池自动分析，生成对应基金池表格 ref:test_fof_fund_pool.py
        type_fundpool = request.POST.get("type_fundpool","")   
        date_fundpool = request.POST.get("date_fundpool","")   
        ### if_purchase 判断基金是否可申购，规避定期开放等品种 | if_purchase=0 对应无限制，1对应可申购
        if_purchase = int( request.POST.get("if_purchase","") ) 
        
        print("Debug- if_purchase= ",if_purchase )
        from funds import fundpool
        class_fundpool = fundpool()
        obj_f= {} 
        obj_f["date_fundpool"] = date_fundpool 
        obj_f["type_fundpool"] = type_fundpool 
        obj_f["if_purchase"] = if_purchase 
        obj_f = class_fundpool.cal_fundpool( obj_f )  

    ###########################################################################
    ### 2,基金收益率排名计算：年初至今
    if "input_fund_ret_rank" in request.POST.keys():  
        ###  
        ret_fund_rank = request.POST.get("ret_fund_rank","")   
        date_fund_rank = request.POST.get("date_fund_rank","")   
        type_fundpool_rank = request.POST.get("type_fundpool_rank","")   
        ### 以防有其他2个同类基金的收益率要算
        ret_fund_rank2 = request.POST.get("ret_fund_rank2","-2")   
        ret_fund_rank3 = request.POST.get("ret_fund_rank3","-2")   
        ##########################################
        ### 导入配置文件
        from config_data import config_data
        config_data_1 = config_data() 

        ### 读取基金收益率文件 || FF-基金研究-主动股票-220520.xlsx
        if int(date_fund_rank) < 20230400 :
            temp_path = config_data_1.obj_config["dict"]["path_wind_terminal"]
        else :
            temp_path= config_data_1.obj_config["dict"]["path_wind_terminal"] + "FF-基金研究\\"
        
        file_name = "FF-基金研究-" + type_fundpool_rank + "-开放申购-" + str(date_fund_rank)[-6:] +".xlsx"
        if not os.path.exists( temp_path + file_name ) :
            file_name = "FF-基金研究-" + type_fundpool_rank + "-" + str(date_fund_rank)[-6:] +".xlsx"
        print("Debug:",file_name )
        
        ##########################################
        ### 
        df_fundpool = pd.read_excel( temp_path + file_name,sheet_name="概况" )  
        ### notes:部分基金收益率存在 "--"
        df_fundpool = df_fundpool[ df_fundpool["收益率（%）"] != "--" ]
        df_fundpool["ret"] = df_fundpool["收益率（%）"]
        df_fundpool["ret"] =df_fundpool["ret"].astype("float")
        
        print("Debug:",df_fundpool.tail() )
        ### calculation
        
        temp_num = df_fundpool[ df_fundpool["ret"] > float(ret_fund_rank ) ]["ret"].count()
        num_all = df_fundpool["ret"].count() 
        ### output 
        rank_str = str(temp_num)+"/"+ str(num_all )
        rank_pct  = round(temp_num/num_all *100,1)
        context["rank_str"] = rank_str
        context["rank_pct"] = rank_pct
        ##########################################
        ### 计算第2、3个基金收益率
        if float(ret_fund_rank2 ) > -100 : 
            temp_num = df_fundpool[ df_fundpool["ret"] > float(ret_fund_rank2 ) ]["ret"].count()
            num_all = df_fundpool["ret"].count() 
            ### output 
            rank_str = str(temp_num)+"/"+ str(num_all )
            rank_pct  = round(temp_num/num_all *100,1)
            context["rank_str2"] = rank_str
            context["rank_pct2"] = rank_pct
    
        if float(ret_fund_rank3 ) > -100 : 
            temp_num = df_fundpool[ df_fundpool["ret"] > float(ret_fund_rank3 ) ]["ret"].count()
            num_all = df_fundpool["ret"].count() 
            ### output 
            rank_str = str(temp_num)+"/"+ str(num_all )
            rank_pct  = round(temp_num/num_all *100,1)
            context["rank_str3"] = rank_str
            context["rank_pct3"] = rank_pct
    
    ###########################################################################
    ### 3,基金收益率排名计算：周/月基金分类统计top5
    if "input_fund_chg_stat" in request.POST.keys():  
        ### 根据基金中短期业绩表，统计不同区间业绩
        ### notes:暂时js无法实现表格任意一列的排序
        date_fund_stat = request.POST.get("date_fund_stat","")   
        fund_type = request.POST.get("fund_type","")   
        from funds import fund_exhi
        class_fund_exhi = fund_exhi()    
        obj_f = {} 
        obj_f["date_fund_stat"] = date_fund_stat
        obj_f["fund_type"] = fund_type
        obj_f = class_fund_exhi.stat_fund_chg_pct( obj_f ) 
        ### output :obj_f["df_fund"] , obj_f["df_fund_sub"] 
        ### 只取前100名
        context["df_fund_exhi"] = obj_f["df_fund_exhi"].T

    return render(request,"ciss_exhi/fund_fof/fund_data_manage.html",context) 




######################################################################################################
######################################################################################################
@csrf_protect
@requires_csrf_token
@csrf_exempt
def fund_indi_template(request): 
    #################################################################################
    ### 基金分析指标和基金入池模板等
    ### 定义dict对象 context
    print("request.POST", type(request.POST.keys() ),request.POST.keys() )
    context={"info":"none"} 
    ### 时间
    context["time_now_str"] = time_now_str
    context["time_now_str_pre1d"] =time_now_str_pre1d  
    
    ### 导入常用地址信息
    import os  
    # sys.path.append(path_ciss_rc+ "config\\")
    from config_data import config_data
    config_db_1 = config_data() 
    path_ciss_rc = config_db_1.obj_config["dict"]["path_ciss_rc"]
    path_ciss_web = config_db_1.obj_config["dict"]["path_ciss_web"] 
    path_fund = config_db_1.obj_config["dict"]["path_fundpool"]
    path_fund_indi = config_db_1.obj_config["dict"]["path_fund_indi"] 
    path_data_pms = config_db_1.obj_config["dict"]["path_data_pms"] 

    ###########################################################################
    ### 对应html， fund_indi_template.html
    ###########################################################################
    ### 更新日期数据至未来6个月
    if "input_check_data_update_date" in request.POST.keys():  
        ### 引用times.py 里的 manage_date_trade
        from times import times
        times1 = times()
        obj_dt = times1.manage_date_trade()

        ### 
        import datetime as dt  
        date_latest = obj_dt["date_latest"]
        if type( date_latest ) == dt.datetime : 
            date_latest = dt.datetime.strftime( date_latest, "%Y%m%d")
        date_end = obj_dt["date_end"]

        
        context["output_latest_date"] = [date_latest, date_end] 

    ###########################################################################
    ### 新增月末跟踪的债券基金代码列表 
    if "input_get_fundlist_bond_1m" in request.POST.keys():  
        ### 月末日期
        date_get_fundlist =  request.POST.get("date_get_fundlist_1m","")  
        print("len =",len( date_get_fundlist ) )
        # if len( date_get_fundlist ) != 8 :

        #     ### 引用times.py 里的 manage_date_trade
        #     from times import times
        #     times1 = times()
        #     obj_dt = times1.manage_date_trade() 
        #     date_get_fundlist = obj_dt["date_latest_str"]  

        ###########################################################################
        ### 导入新增的代码列表 | 1，定期新增：每个月末，从几个基金策略中纳入的品种；2，不定期新增，临时决定新增的基金品种；
        ### input：sheet=fundcode,file=fund_indi_manage.xlsx ； path=C:\rc_2023\rc_202X\data_pms\
        df_fundcode = pd.read_excel(path_data_pms+ "fund_indi_manage.xlsx",sheet_name= "fundcode" )
        num_fund_code = len( df_fundcode.index )

        ###########################################################################
        ### 3，存量评估和延续：对近6个月的绩效进行评估，总分最差的后20%剔除；对特别品种可以标记不剔除；
        ### 导入新增的代码列表 | 存量评估和延续：对近6个月的绩效进行评估，总分最差的后20%剔除；对特别品种可以标记不剔除；
        file_fundpool = "fundcode_bond_" + date_get_fundlist + ".xlsx"
        
        ### 判断文件是否存在，如果存在要合并数据
        if os.path.exists( path_fund_indi+ file_fundpool ) :
            ### 读取已有的数据
            df_orignal = pd.read_excel( path_fund_indi+ file_fundpool )
            if not "code" in df_fundcode.columns :
                df_fundcode["code"] = df_fundcode["基金代码"]
            ###  合并基金代码列表
            for temp_i in df_fundcode.index : 
                for temp_col in df_fundcode.columns :  
                    df_orignal.loc[temp_i,temp_col ] =  df_fundcode.loc[temp_i, temp_col] 

        else : 
            df_orignal = df_fundcode

        ###########################################################################
        ### TODO 对近6个月的绩效进行评估，总分最差的后20%剔除；对特别品种可以标记不剔除；
        ### column=["score_pre_1m","_pre_2m","_pre_3m","_pre_4m","_pre_5m","_pre_6m"]
        # 如果前边没有5个月的数据怎么办？
        # score_pre_1m 的计算：导入上个月基金指标文件，如 file= fund_indi_20230331.xlsx 
        # 基金指标：近1个月收益率、近1个月最大回撤、alphs，相对收益率、相对回撤， 加权打分 
        # notes: fundcode_bond_20230331.xlsx 中有基金代码列表 

        ###########################################################################
        ### 合并数据后，剔除重复项
        df_orignal = df_orignal.drop_duplicates( keep="first") 
        
        ### 保存至excel
        ### Qs：用pd.ExcelWriter导出excel时，给定sheet名称如果已经存在，Sheet1 会变成 Sheet11
        ### pd.ExcelWriter 只能操作已经有的excel文件
        ### *mode='a'代表append，可以实现追加sheet；如果不写mode,则默认mode='w'，会把原有的sheet覆盖。
        # writer = pd.ExcelWriter( path_fund_indi + file_output ,mode='a',engine='openpyxl') 
        # df_orignal.to_excel( writer ,sheet_name="Sheet1",header=True,index=False )  
        # writer.save()
        # writer.close()
        ### 
        df_orignal.to_excel(path_fund_indi + file_fundpool ,sheet_name="Sheet1",header=True,index=False )  

        ### output
        context["date_month_end"] = [date_get_fundlist] 
        context["num_fund_code"] = [num_fund_code] 

    ###########################################################################
    ### 获取和计算债券类基金指标数据
    ###########################################################################
    ### 方式一：单月末+单个基金代码
    if "input_fund_indi_1month_1fund" in request.POST.keys():  
        ### 根据输入的“基金代码、时间等参数、基金和持仓指标”，获取、计算标准化的基金数据。只提取月末或者季末数据。
        ### 指标类型：基本信息、绩效分析、持仓分析、基准数据。
        ### date_fund_indi 默认格式 20230413
        date_fund_indi = request.POST.get("date_fund_indi_1m1f","")   
        code_fund = request.POST.get("code_fund_indi_get_data","")   

        ################################################
        obj_fund = {} 
        ## "20230331" 
        obj_fund["date_m_end"] = date_fund_indi 
        
        ################################################
        ### 获取基金指标：基本信息、基准指数信息，和绩效分析、持仓分析、基准数据。
        # 需要在 "fund_indi_manage.xlsx"设置好 基金代码、基准指数代码、指标。
        from fund_indicator_windapi import fund_indicator
        class_fund_indicator = fund_indicator()    
        
        ### manage_type=1，给定基金列表
        obj_fund["list_code"] = [ code_fund ]
        obj_fund["manage_type"] = "1"
        obj_fund = class_fund_indicator.manage_fund_indi(obj_fund)    
        ### obj_fund["df_index"] 

        context["output_fund_indi_get_data_1m1f"] = obj_fund["df_index"].head(2).T

    ###########################################################################
    ### 方式二：单月末+基金代码列表 | 给定基金列表
    if "input_fund_indi_1month_Mfund" in request.POST.keys():  
        ### 导入基金代码的月末时间
        date_fundcode_input   = request.POST.get("date_fundcode_input_1mmf","")    
        ### 基金指标对应的月末时间
        date_fund_indi_output   = request.POST.get("date_fund_indi_output_1mmf","")    

        ################################################
        obj_fund = {}  
        obj_fund["date_m_end"] = date_fund_indi_output 
        
        ################################################
        ### 获取基金指标：基本信息、基准指数信息，和绩效分析、持仓分析、基准数据。  
        from fund_indicator_windapi import fund_indicator
        class_fund_indicator = fund_indicator()    
        
        ### 基金列表:sheet=fundcode,file=fund_indi_manage.xlsx,path=C:\rc_2023\rc_202X\data_pms 
#       ##########################################
        ### fund code from excel sheet || file=fundcode_bond_20220531.xlsx ； path=C:\rc_2023\rc_202X\data_pms\fund\fund_indi\
        temp_file = "fundcode_bond_" + date_fundcode_input +".xlsx" 
        df_fundcode = pd.read_excel( path_fund_indi + temp_file  )
        if "code" in df_fundcode.columns :
            list_code = list( df_fundcode["code"] )
        elif "基金代码" in df_fundcode.columns :
            list_code = list( df_fundcode["基金代码"] )

        ### 去除 list_code 中的重复项
        list_code =  list(set( list_code ))


        obj_fund["list_code"] = list_code
        obj_fund["manage_type"] = "1"
        obj_fund = class_fund_indicator.manage_fund_indi(obj_fund)   

        ################################################
        ### 展示数据   obj_fund["df_fund_indi"]
        ### 导出数据

        context["output_fund_indi_get_data_1mmf"] = obj_fund["df_fund_indi"].head(5).T

    ###########################################################################
    ### 方式三：日期区间+基金代码列表 | 给定基金列表 TODO
    # input_fund_indi_Nmonth_Mfund

    ###########################################################################
    ### 生成Excel/Word基金入池研究报告
    if "input_gen_excel_word_file_fund" in request.POST.keys():  
        ### 根据输入的“基金代码、时间等参数、基金和持仓指标”，获取、计算标准化的基金数据。只提取月末或者季末数据。
        ### 指标类型：基本信息、绩效分析、持仓分析、基准数据。
        date_gen = request.POST.get("date_gen_excel_word_file_fund","")   


        ################################################
        ### 获取word表格内容
        # C:\rc_2023\rc_202X\data_pms\\fund\\
        temp_path = path_fund +"fund_report\\"
        file_name = "基金入池研究报告.docx"

        import docx
        doc= docx.Document( temp_path+ file_name )
        ### doc.tables[0].cell(0,0).text = '证券信息'

        ################################################
        ### 第一行，基本信息
        # doc.tables[0].cell(1,1).text = 证券信息
        print("DEBUG= doc.tables[0].cell(1,1).text=" ,doc.tables[0].cell(0,0).text  )
        doc.tables[0].cell(1,1).text = "基金打分和调研数据表管理"

        # doc.tables[0].cell(1,2).text 基金名称

        # doc.tables[0].cell(2,2).text 申请加入固定收益类基金池名称

        # doc.tables[0].cell(2,3).text 场内/场外

        ################################################
        ### 第二行，申请理由
        # doc.tables[0].cell(2,1).text


        ################################################
        ### 第三行，满足条件
        # doc.tables[0].cell(3,1).text

        ################################################################################################
        ### 第二个表格
        ### doc.tables[1].cell(0,0).text = '基金研究报告'








        ################################################
        ### output file
        file_output = "基金入池研究报告_"+ date_gen +".docx"
        doc.save( temp_path+ file_output )

























    return render(request,"ciss_exhi/fund_fof/fund_indi_template.html",context) 

    

























